# -*- coding: utf-8 -*-
from bs4 import BeautifulSoup
import requests
import os.path
from docx import Document


def getPageData(url):
    r = requests.get(url)
    data = r.text
    soup = BeautifulSoup(data, 'lxml')
    title = getTitle(soup)
    lyrics = getLyrics(soup)
    return title, lyrics


def getTitle(soup):
    return soup.find('h1', class_='page-title').text.encode('utf-8').strip()


def getLyrics(soup):
    return soup.find('div', {'class':'content-text-inner'}).text.encode('utf-8').strip()


def writeToFile(artist, url, counter):
    title, title_, lyrics = getTitleAndLyrics(url)

    pathToCorpus = './../corpi/{}/raw/'.format(artist)


    if not os.path.exists(pathToCorpus):
        os.makedirs(pathToCorpus)

    fileName = '{}_{}_{}.txt'.format(artist,counter,title_)
    corpusFile= "{}/{}".format(pathToCorpus,fileName)
    try:
        with open(corpusFile, 'a') as theFile:

            theFile.write( title  )
            theFile.write( '\n' )
            theFile.write( '\n' )
            theFile.write( lyrics )
            theFile.write( '\n' )
            theFile.write( '\n' )
            theFile.write( '.' )
            theFile.write( '\n' )
            theFile.write( '\n' )
    except Exception as writeEx:
        print writeEx


def getTitleAndLyrics(url):
    title, lyrics = getPageData(url)
    title_ = title.replace('lyrics', '').strip()
    title_ = title_.replace(' ', '_').strip()
    title_ = title_.replace(',', '').strip()
    return title, title_, lyrics


def writeToDocX(artist, url, counter):
    title, title_, lyrics = getTitleAndLyrics(url)

    pathToCorpus = './../corpi/allSongs/docx/'


    if not os.path.exists(pathToCorpus):
        os.makedirs(pathToCorpus)

    fileName = '{}_{}_{}.docx'.format(artist,counter,title_)
    corpusFile= "{}/{}".format(pathToCorpus,fileName)

    try:

        document = Document()
        core_properties = document.core_properties
        core_properties.author = artist
        document.add_heading(title)
        document.add_paragraph(unicode(lyrics))
        document.save(corpusFile)
    except Exception as docex:
        print "docex error {} {}".format(artist, title)
        print docex





def parseArtist(artists):
    for artist in artists:

        lines=[]
        with open('../links/{}.urls'.format(artist)) as file:
            lines = [line.strip() for line in file]

        print artist, ' has this many song link:', len(lines)

        counter=0
        for line in lines:
            counter = counter + 1
            writeToFile(artist , line, counter)
            writeToDocX(artist , line, counter)

        print "end ", artist



if __name__ == '__main__':
    artists=['pink_floyd','michael_jackson','prince','led_zeppelin']
    parseArtist(artists)







